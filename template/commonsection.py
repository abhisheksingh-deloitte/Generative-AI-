import openai
import openpyxl
from docx import Document


def fetch_data_and_write_to_doc(sheet_name, headings, column_names, output_doc,contents):
    # Load the Excel file
    workbook = openpyxl.load_workbook('../Utils/excel.xlsx')

    # Select the specified sheet
    sheet = workbook[sheet_name]

    # Find the column indices based on the column names
    column_indices = {}
    for heading, column_name in zip(headings, column_names):
        column_index = None
        for cell in sheet[1]:
            if cell.value == column_name:
                column_index = cell.column
                break
        if column_index is None:
            print(f"Column '{column_name}' not found.")
            return
        column_indices[heading] = column_index

    # Create a new document
    doc = Document()

    # Write the headings to the document
    for heading in headings:
        doc.add_paragraph(heading)

    # Fetch the data from the specified columns and write it to the document in front of the respective headings
    num_rows = sheet.max_row
    doc.add_heading(f"{heading}", level=1)
    for row_index in range(2, num_rows + 1):
        values = {}  # Dictionary to store the column values
        for heading, column_index in column_indices.items():
            column_value = sheet.cell(row=row_index, column=column_index).value
            values[heading] = column_value
            # Fetch data from OpenAI
            openai.api_key = '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'
            prompt_values = " ".join([str(value) for value in values.values()])
            prompt = "User Requirement - The functional specification document will specify the details for generating a customer invoice form for Canada. The form type identified is Adobe Forms in BTP, triggered by the business document. This document will outline the necessary information and requirements to create a custom invoice form tailored to the specific needs of Canadian customers."
            spec = f"Choose the content which is applicable for user requirement from below and create the {contents} of a functional specification document in not more than 60 words:"
            response = openai.Completion.create(
                engine='text-davinci-003',
                prompt=prompt + spec + f" {prompt_values}",
                max_tokens=2000
            )
            data = response.choices[0].text.strip()
            doc.add_paragraph(data)

    # Save the document
    doc.save(output_doc)
    print(f"Data written to '{output_doc}' successfully.")


# Example usage
headings = ["Assumptions"]
column_names = ["fs-general_fields-assumptions"]
fetch_data_and_write_to_doc("Sheet1", headings, column_names, "output.docx","assumptions")
